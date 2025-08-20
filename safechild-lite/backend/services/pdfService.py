import os
import io
import logging
import tempfile
from typing import Dict, Any, Optional, List, Union
from pathlib import Path
import json
from datetime import datetime

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Word document generation will be disabled.")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab not available. PDF generation will be disabled.")

from ..models.complaintModel import ComplaintData, ComplaintStatus, PriorityLevel
from ..utils.textCleaner import TextCleaner
from ..utils.timeUtils import TimeUtils

logger = logging.getLogger(__name__)

class PDFService:
    """Service for generating PDF and Word documents"""
    
    def __init__(self):
        self.text_cleaner = TextCleaner()
        self.time_utils = TimeUtils()
        
        # Document templates directory
        self.templates_dir = Path(os.getenv("PDF_TEMPLATES_DIR", "templates/documents"))
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        
        # Output directory
        self.output_dir = Path(os.getenv("PDF_OUTPUT_DIR", "temp/documents"))
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Document settings
        self.default_page_size = os.getenv("PDF_DEFAULT_PAGE_SIZE", "A4")
        self.default_font_size = int(os.getenv("PDF_DEFAULT_FONT_SIZE", "12"))
        self.default_margin = float(os.getenv("PDF_DEFAULT_MARGIN", "1.0"))
        
        # Check library availability
        self.docx_available = DOCX_AVAILABLE
        self.reportlab_available = REPORTLAB_AVAILABLE
        
        if not self.docx_available and not self.reportlab_available:
            logger.warning("No document generation libraries available. PDF service will be limited.")
        
        logger.info(f"PDF Service initialized. DOCX: {self.docx_available}, ReportLab: {self.reportlab_available}")
    
    async def generate_complaint_document(
        self, 
        complaint_data: ComplaintData, 
        document_type: str = "pdf",
        template_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """Generate complaint document in specified format"""
        try:
            if document_type.lower() == "pdf":
                if not self.reportlab_available:
                    return {
                        "success": False,
                        "error": "PDF generation not available. ReportLab library not installed.",
                        "document_data": None
                    }
                return await self._generate_pdf_complaint(complaint_data, template_name)
            
            elif document_type.lower() in ["docx", "word", "doc"]:
                if not self.docx_available:
                    return {
                        "success": False,
                        "error": "Word document generation not available. python-docx library not installed.",
                        "document_data": None
                    }
                return await self._generate_docx_complaint(complaint_data, template_name)
            
            else:
                return {
                    "success": False,
                    "error": f"Unsupported document type: {document_type}",
                    "supported_types": ["pdf", "docx", "word", "doc"],
                    "document_data": None
                }
                
        except Exception as e:
            logger.error(f"Error generating complaint document: {e}")
            return {
                "success": False,
                "error": f"Document generation failed: {str(e)}",
                "document_data": None
            }
    
    async def generate_safety_report(
        self, 
        report_data: Dict[str, Any], 
        document_type: str = "pdf",
        include_charts: bool = False
    ) -> Dict[str, Any]:
        """Generate safety report document"""
        try:
            if document_type.lower() == "pdf":
                if not self.reportlab_available:
                    return {
                        "success": False,
                        "error": "PDF generation not available",
                        "document_data": None
                    }
                return await self._generate_pdf_safety_report(report_data, include_charts)
            
            elif document_type.lower() in ["docx", "word"]:
                if not self.docx_available:
                    return {
                        "success": False,
                        "error": "Word document generation not available",
                        "document_data": None
                    }
                return await self._generate_docx_safety_report(report_data, include_charts)
            
            else:
                return {
                    "success": False,
                    "error": f"Unsupported document type: {document_type}",
                    "document_data": None
                }
                
        except Exception as e:
            logger.error(f"Error generating safety report: {e}")
            return {
                "success": False,
                "error": f"Report generation failed: {str(e)}",
                "document_data": None
            }
    
    async def generate_emergency_alert_report(
        self, 
        alert_data: Dict[str, Any], 
        document_type: str = "pdf"
    ) -> Dict[str, Any]:
        """Generate emergency alert report"""
        try:
            if document_type.lower() == "pdf":
                if not self.reportlab_available:
                    return {
                        "success": False,
                        "error": "PDF generation not available",
                        "document_data": None
                    }
                return await self._generate_pdf_emergency_report(alert_data)
            
            elif document_type.lower() in ["docx", "word"]:
                if not self.docx_available:
                    return {
                        "success": False,
                        "error": "Word document generation not available",
                        "document_data": None
                    }
                return await self._generate_docx_emergency_report(alert_data)
            
            else:
                return {
                    "success": False,
                    "error": f"Unsupported document type: {document_type}",
                    "document_data": None
                }
                
        except Exception as e:
            logger.error(f"Error generating emergency report: {e}")
            return {
                "success": False,
                "error": f"Emergency report generation failed: {str(e)}",
                "document_data": None
            }
    
    async def batch_generate_documents(
        self, 
        documents: List[Dict[str, Any]], 
        document_type: str = "pdf"
    ) -> Dict[str, Any]:
        """Generate multiple documents in batch"""
        try:
            if not documents or len(documents) > 20:  # Limit batch size
                return {
                    "success": False,
                    "error": "Invalid batch size. Must be between 1 and 20 documents.",
                    "results": []
                }
            
            results = []
            for i, doc_data in enumerate(documents):
                doc_type = doc_data.get("type", "complaint")
                
                if doc_type == "complaint":
                    result = await self.generate_complaint_document(
                        doc_data["data"], 
                        document_type
                    )
                elif doc_type == "safety_report":
                    result = await self.generate_safety_report(
                        doc_data["data"], 
                        document_type
                    )
                elif doc_type == "emergency_report":
                    result = await self.generate_emergency_alert_report(
                        doc_data["data"], 
                        document_type
                    )
                else:
                    result = {
                        "success": False,
                        "error": f"Unknown document type: {doc_type}",
                        "document_data": None
                    }
                
                result["index"] = i
                result["document_type"] = doc_type
                results.append(result)
            
            success_count = sum(1 for r in results if r["success"])
            
            return {
                "success": success_count > 0,
                "total_documents": len(documents),
                "successful_generations": success_count,
                "failed_generations": len(documents) - success_count,
                "results": results,
                "timestamp": self.time_utils.get_current_timestamp()
            }
            
        except Exception as e:
            logger.error(f"Error in batch document generation: {e}")
            return {
                "success": False,
                "error": f"Batch generation failed: {str(e)}",
                "results": []
            }
    
    async def _generate_pdf_complaint(
        self, 
        complaint_data: ComplaintData, 
        template_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """Generate PDF complaint document using ReportLab"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_file:
                temp_path = temp_file.name
            
            # Create PDF document
            doc = SimpleDocTemplate(
                temp_path,
                pagesize=A4,
                rightMargin=self.default_margin * inch,
                leftMargin=self.default_margin * inch,
                topMargin=self.default_margin * inch,
                bottomMargin=self.default_margin * inch
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=20,
                alignment=TA_CENTER,
                textColor=colors.darkblue
            )
            
            # Build story (content)
            story = []
            
            # Title
            story.append(Paragraph("CHILD SAFETY INCIDENT COMPLAINT", title_style))
            story.append(Spacer(1, 20))
            
            # Basic Information
            story.append(Paragraph("Basic Information", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            basic_info = [
                ["Complaint ID:", complaint_data.complaint_id or "N/A"],
                ["Date Filed:", complaint_data.date_filed or "N/A"],
                ["Status:", complaint_data.status.value if complaint_data.status else "N/A"],
                ["Priority:", complaint_data.priority.value if complaint_data.priority else "N/A"]
            ]
            
            basic_table = Table(basic_info, colWidths=[2*inch, 4*inch])
            basic_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(basic_table)
            story.append(Spacer(1, 20))
            
            # Child Information
            if complaint_data.child_information:
                story.append(Paragraph("Child Information", styles['Heading2']))
                story.append(Spacer(1, 12))
                
                child_info = [
                    ["Name:", complaint_data.child_information.get("name", "N/A")],
                    ["Age:", str(complaint_data.child_information.get("age", "N/A"))],
                    ["Gender:", complaint_data.child_information.get("gender", "N/A")],
                    ["School/Institution:", complaint_data.child_information.get("school", "N/A")]
                ]
                
                child_table = Table(child_info, colWidths=[2*inch, 4*inch])
                child_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(child_table)
                story.append(Spacer(1, 20))
            
            # Incident Details
            story.append(Paragraph("Incident Details", styles['Heading2']))
            story.append(Spacer(1, 12))
            
            if complaint_data.incident_details:
                incident_text = complaint_data.incident_details.get("description", "No description provided")
                story.append(Paragraph(incident_text, styles['Normal']))
                story.append(Spacer(1, 12))
                
                # Incident metadata
                incident_meta = [
                    ["Incident Type:", complaint_data.incident_details.get("type", "N/A")],
                    ["Date of Incident:", complaint_data.incident_details.get("date", "N/A")],
                    ["Location:", complaint_data.incident_details.get("location", "N/A")],
                    ["Witnesses:", complaint_data.incident_details.get("witnesses", "None")]
                ]
                
                incident_table = Table(incident_meta, colWidths=[2*inch, 4*inch])
                incident_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(incident_table)
                story.append(Spacer(1, 20))
            
            # Guardian Information
            if complaint_data.guardian_information:
                story.append(Paragraph("Guardian Information", styles['Heading2']))
                story.append(Spacer(1, 12))
                
                guardian_info = [
                    ["Name:", complaint_data.guardian_information.get("name", "N/A")],
                    ["Relationship:", complaint_data.guardian_information.get("relationship", "N/A")],
                    ["Phone:", complaint_data.guardian_information.get("phone", "N/A")],
                    ["Email:", complaint_data.guardian_information.get("email", "N/A")],
                    ["Address:", complaint_data.guardian_information.get("address", "N/A")]
                ]
                
                guardian_table = Table(guardian_info, colWidths=[2*inch, 4*inch])
                guardian_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(guardian_table)
                story.append(Spacer(1, 20))
            
            # Additional Details
            if complaint_data.additional_details:
                story.append(Paragraph("Additional Details", styles['Heading2']))
                story.append(Spacer(1, 12))
                
                additional_text = complaint_data.additional_details.get("notes", "No additional notes")
                story.append(Paragraph(additional_text, styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Footer
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=8,
                alignment=TA_CENTER,
                textColor=colors.grey
            )
            story.append(Paragraph(
                f"Generated on {self.time_utils.get_current_timestamp()} by SafeChild-Lite System",
                footer_style
            ))
            
            # Build PDF
            doc.build(story)
            
            # Read generated PDF
            with open(temp_path, "rb") as f:
                pdf_data = f.read()
            
            # Clean up temporary file
            os.unlink(temp_path)
            
            return {
                "success": True,
                "document_data": pdf_data,
                "document_type": "pdf",
                "file_size": len(pdf_data),
                "filename": f"complaint_{complaint_data.complaint_id or 'unknown'}.pdf",
                "timestamp": self.time_utils.get_current_timestamp()
            }
            
        except Exception as e:
            logger.error(f"Error generating PDF complaint: {e}")
            return {
                "success": False,
                "error": f"PDF generation failed: {str(e)}",
                "document_data": None
            }
    
    async def _generate_docx_complaint(
        self, 
        complaint_data: ComplaintData, 
        template_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """Generate Word document complaint using python-docx"""
        try:
            # Create new document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = "Child Safety Incident Complaint"
            doc.core_properties.author = "SafeChild-Lite System"
            doc.core_properties.subject = "Child Safety Incident Report"
            
            # Title
            title = doc.add_heading("CHILD SAFETY INCIDENT COMPLAINT", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Basic Information
            doc.add_heading("Basic Information", level=1)
            
            # Create table for basic info
            basic_table = doc.add_table(rows=4, cols=2)
            basic_table.style = 'Table Grid'
            
            basic_data = [
                ["Complaint ID:", complaint_data.complaint_id or "N/A"],
                ["Date Filed:", complaint_data.date_filed or "N/A"],
                ["Status:", complaint_data.status.value if complaint_data.status else "N/A"],
                ["Priority:", complaint_data.priority.value if complaint_data.priority else "N/A"]
            ]
            
            for i, (label, value) in enumerate(basic_data):
                basic_table.cell(i, 0).text = label
                basic_table.cell(i, 1).text = value
            
            doc.add_paragraph()  # Add spacing
            
            # Child Information
            if complaint_data.child_information:
                doc.add_heading("Child Information", level=1)
                
                child_table = doc.add_table(rows=4, cols=2)
                child_table.style = 'Table Grid'
                
                child_data = [
                    ["Name:", complaint_data.child_information.get("name", "N/A")],
                    ["Age:", str(complaint_data.child_information.get("age", "N/A"))],
                    ["Gender:", complaint_data.child_information.get("gender", "N/A")],
                    ["School/Institution:", complaint_data.child_information.get("school", "N/A")]
                ]
                
                for i, (label, value) in enumerate(child_data):
                    child_table.cell(i, 0).text = label
                    child_table.cell(i, 1).text = value
                
                doc.add_paragraph()  # Add spacing
            
            # Incident Details
            doc.add_heading("Incident Details", level=1)
            
            if complaint_data.incident_details:
                incident_desc = complaint_data.incident_details.get("description", "No description provided")
                doc.add_paragraph(incident_desc)
                
                # Incident metadata table
                incident_table = doc.add_table(rows=4, cols=2)
                incident_table.style = 'Table Grid'
                
                incident_data = [
                    ["Incident Type:", complaint_data.incident_details.get("type", "N/A")],
                    ["Date of Incident:", complaint_data.incident_details.get("date", "N/A")],
                    ["Location:", complaint_data.incident_details.get("location", "N/A")],
                    ["Witnesses:", complaint_data.incident_details.get("witnesses", "None")]
                ]
                
                for i, (label, value) in enumerate(incident_data):
                    incident_table.cell(i, 0).text = label
                    incident_table.cell(i, 1).text = value
                
                doc.add_paragraph()  # Add spacing
            
            # Guardian Information
            if complaint_data.guardian_information:
                doc.add_heading("Guardian Information", level=1)
                
                guardian_table = doc.add_table(rows=5, cols=2)
                guardian_table.style = 'Table Grid'
                
                guardian_data = [
                    ["Name:", complaint_data.guardian_information.get("name", "N/A")],
                    ["Relationship:", complaint_data.guardian_information.get("relationship", "N/A")],
                    ["Phone:", complaint_data.guardian_information.get("phone", "N/A")],
                    ["Email:", complaint_data.guardian_information.get("email", "N/A")],
                    ["Address:", complaint_data.guardian_information.get("address", "N/A")]
                ]
                
                for i, (label, value) in enumerate(guardian_data):
                    guardian_table.cell(i, 0).text = label
                    guardian_table.cell(i, 1).text = value
                
                doc.add_paragraph()  # Add spacing
            
            # Additional Details
            if complaint_data.additional_details:
                doc.add_heading("Additional Details", level=1)
                additional_notes = complaint_data.additional_details.get("notes", "No additional notes")
                doc.add_paragraph(additional_notes)
            
            # Footer
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Generated on {self.time_utils.get_current_timestamp()} by SafeChild-Lite System"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_path = temp_file.name
            
            doc.save(temp_path)
            
            # Read generated document
            with open(temp_path, "rb") as f:
                docx_data = f.read()
            
            # Clean up temporary file
            os.unlink(temp_path)
            
            return {
                "success": True,
                "document_data": docx_data,
                "document_type": "docx",
                "file_size": len(docx_data),
                "filename": f"complaint_{complaint_data.complaint_id or 'unknown'}.docx",
                "timestamp": self.time_utils.get_current_timestamp()
            }
            
        except Exception as e:
            logger.error(f"Error generating DOCX complaint: {e}")
            return {
                "success": False,
                "error": f"DOCX generation failed: {str(e)}",
                "document_data": None
            }
    
    async def _generate_pdf_safety_report(
        self, 
        report_data: Dict[str, Any], 
        include_charts: bool = False
    ) -> Dict[str, Any]:
        """Generate PDF safety report"""
        # Implementation for safety report PDF generation
        # This would be similar to complaint PDF but with different content structure
        pass
    
    async def _generate_docx_safety_report(
        self, 
        report_data: Dict[str, Any], 
        include_charts: bool = False
    ) -> Dict[str, Any]:
        """Generate Word document safety report"""
        # Implementation for safety report DOCX generation
        pass
    
    async def _generate_pdf_emergency_report(
        self, 
        alert_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate PDF emergency report"""
        # Implementation for emergency report PDF generation
        pass
    
    async def _generate_docx_emergency_report(
        self, 
        alert_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate Word document emergency report"""
        # Implementation for emergency report DOCX generation
        pass
    
    async def health_check(self) -> Dict[str, Any]:
        """Check PDF service health"""
        try:
            return {
                "status": "healthy",
                "service": "PDF Service",
                "docx_available": self.docx_available,
                "reportlab_available": self.reportlab_available,
                "templates_directory": str(self.templates_dir),
                "output_directory": str(self.output_dir),
                "default_page_size": self.default_page_size,
                "default_font_size": self.default_font_size,
                "timestamp": self.time_utils.get_current_timestamp()
            }
        except Exception as e:
            logger.error(f"PDF service health check failed: {e}")
            return {
                "status": "unhealthy",
                "service": "PDF Service",
                "error": str(e),
                "timestamp": self.time_utils.get_current_timestamp()
            }
    
    def get_usage_stats(self) -> Dict[str, Any]:
        """Get PDF service usage statistics"""
        return {
            "docx_available": self.docx_available,
            "reportlab_available": self.reportlab_available,
            "templates_directory": str(self.templates_dir),
            "output_directory": str(self.output_dir),
            "default_page_size": self.default_page_size,
            "default_font_size": self.default_font_size,
            "default_margin": self.default_margin
        }
